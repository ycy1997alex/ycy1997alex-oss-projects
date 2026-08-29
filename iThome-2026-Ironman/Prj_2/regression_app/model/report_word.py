"""Word 分析報告輸出（python-docx）。

報告的定位是「這份分析做了什麼、結果怎麼讀」，不是資料傾印 ——
逐列的數字在 Predictions.xlsx，這裡只放結論需要的表與圖。
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from regression_app.model.bolasso import BolassoResult
from regression_app.model.modeling import METRIC_ORDER, ModelResult
from regression_app.model.schema import Dataset

REPORT_NAME = "Analysis_Report.docx"

_BODY_FONT = "Microsoft JhengHei"
_FIGURE_WIDTH = Inches(6.2)
_HEADING_COLOR = RGBColor(0x1A, 0x5F, 0xA8)


def write_report(
    dataset: Dataset,
    models: list[ModelResult],
    bolassos: dict[str, BolassoResult],
    figures: dict[str, Path],
    out_dir: Path,
    input_names: list[str],
    app_version: str,
) -> Path:
    """產生 Analysis_Report.docx。figures 是圖表用途 -> 檔案路徑的對照表。"""
    doc = Document()
    _setup_styles(doc)

    _title_block(doc, dataset, models, input_names, app_version)
    _method_section(doc, bolassos, models)
    _data_section(doc, dataset, figures, input_names)
    _per_target_sections(doc, models, bolassos, figures)
    _comparison_section(doc, models, figures)
    _caveats_section(doc, models, bolassos)

    path = out_dir / REPORT_NAME
    doc.save(path)
    return path


# --------------------------------------------------------------------------
# 版面基礎
# --------------------------------------------------------------------------

def _setup_styles(doc: Document) -> None:
    """設定中文字型。只改 latin 的 font.name 不夠，東亞字型要另外指定。"""
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), _BODY_FONT)

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = _BODY_FONT
        style.font.color.rgb = _HEADING_COLOR
        style.element.rPr.rFonts.set(qn("w:eastAsia"), _BODY_FONT)


def _fmt(value: object, digits: int = 4) -> str:
    """數值統一格式化；p 值太小就寫 < 0.001，不要印出一串 0。"""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "—"
    if isinstance(value, (int,)) and not isinstance(value, bool):
        return f"{value:,}"
    if isinstance(value, float):
        if value != 0 and abs(value) < 10 ** (-digits):
            return f"< {10 ** (-digits):g}"
        return f"{value:,.{digits}f}"
    return str(value)


def _add_table(doc: Document, frame: pd.DataFrame, digits: int = 4) -> None:
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"

    for cell, name in zip(table.rows[0].cells, frame.columns):
        cell.text = str(name)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)

    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = _fmt(value, digits)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


def _add_figure(doc: Document, path: Path | None, caption: str) -> None:
    if path is None or not Path(path).is_file():
        return
    doc.add_picture(str(path), width=_FIGURE_WIDTH)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph(caption)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7E, 0x80, 0x81)


# --------------------------------------------------------------------------
# 各章節
# --------------------------------------------------------------------------

def _title_block(
    doc: Document, dataset: Dataset, models: list[ModelResult], input_names: list[str], app_version: str
) -> None:
    heading = doc.add_heading("迴歸統計分析報告", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = pd.DataFrame(
        [
            {"項目": "來源檔案", "內容": dataset.path.name},
            {"項目": "產生時間", "內容": datetime.now().strftime("%Y-%m-%d %H:%M:%S")},
            {"項目": "樣本數", "內容": f"{dataset.n_rows:,} 筆"},
            {"項目": "納入分析的輸入變項", "內容": f"{len(input_names)} 個：" + "、".join(input_names)},
            {"項目": "結果變項", "內容": f"{len(models)} 個：" + "、".join(m.target for m in models)},
            {"項目": "產生工具", "內容": f"迴歸統計分析工具 v{app_version}"},
        ]
    )
    _add_table(doc, meta)


def _method_section(doc: Document, bolassos: dict[str, BolassoResult], models: list[ModelResult]) -> None:
    doc.add_heading("1. 分析方法", level=1)
    sample = next(iter(bolassos.values()))

    doc.add_paragraph(
        f"變項篩選採用 Bolasso（Bootstrap-enhanced Lasso）。做法是對原始樣本做 "
        f"{sample.n_bootstrap} 次可放回抽樣，每一次都配適一個 Lasso 迴歸，記錄每個候選變項的係數是否被壓成 0；"
        f"在夠高比例的抽樣中都存活下來的變項，才視為穩定有貢獻。"
    )
    doc.add_paragraph(
        f"懲罰強度 α 先以全樣本的 5-fold LassoCV 選定一次後固定，讓 {sample.n_bootstrap} 次抽樣的懲罰一致，"
        f"選入頻率之間才能直接比較。選入門檻為 {sample.threshold:.0%}。"
    )
    doc.add_paragraph(
        "最終模型以選入的變項配適普通最小平方法（OLS）迴歸，取得係數、標準誤、t 值、p 值與 95% 信賴區間。"
        "評估指標同時報告樣本內與 5-fold 交叉驗證兩組數字：樣本內指標必然偏樂觀，"
        "交叉驗證的數字才接近模型對新資料的實際表現。"
    )


def _data_section(doc: Document, dataset: Dataset, figures: dict[str, Path], input_names: list[str]) -> None:
    doc.add_heading("2. 資料概況", level=1)

    doc.add_heading("2.1 欄位角色", level=2)
    table = pd.DataFrame(
        [
            {
                "欄位": c.name,
                "角色": c.role_raw,
                "資料類型": c.data_type,
                "本次納入": "是" if (c.name in input_names or c.role == "Result") else "否",
                "備註": c.note or "—",
            }
            for c in dataset.columns
        ]
    )
    _add_table(doc, table)

    doc.add_heading("2.2 分布與相關", level=2)
    _add_figure(doc, figures.get("violin"), "圖 2-1　連續型欄位分布（小提琴圖，內含盒鬚圖）")
    _add_figure(doc, figures.get("correlation"), "圖 2-2　Pearson 與 Spearman 相關矩陣")
    _add_figure(doc, figures.get("pairplot"), "圖 2-3　連續型欄位 Pair Plot")

    doc.add_paragraph(
        "個別變項的直方圖存放於 Visualization\\Variable Analysis，"
        "輸入變項對結果變項的散布圖存放於 Visualization\\Input×Result Analysis，"
        "數量較多，未逐張收進本報告。"
    )


def _per_target_sections(
    doc: Document, models: list[ModelResult], bolassos: dict[str, BolassoResult], figures: dict[str, Path]
) -> None:
    doc.add_heading("3. 各結果變項的模型", level=1)

    for i, model in enumerate(models, start=1):
        bol = bolassos[model.target]
        doc.add_heading(f"3.{i} {model.target}", level=2)

        doc.add_heading(f"3.{i}.1 變項篩選", level=3)
        freq_table = pd.DataFrame(
            {
                "候選變項": bol.frequency.index,
                "選入頻率": [f"{v:.1%}" for v in bol.frequency.to_numpy()],
                "是否選入": ["是" if f in set(bol.selected) else "否" for f in bol.frequency.index],
            }
        )
        _add_table(doc, freq_table)
        doc.add_paragraph(
            f"α = {bol.alpha:.4f}，{bol.n_bootstrap} 次 bootstrap，門檻 {bol.threshold:.0%}；"
            f"共選入 {bol.n_selected} 個變項：" + "、".join(bol.selected) + "。"
        )
        if bol.fallback_note:
            para = doc.add_paragraph(f"註：{bol.fallback_note}")
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xC4, 0x00, 0x2C)

        doc.add_heading(f"3.{i}.2 迴歸模型", level=3)
        para = doc.add_paragraph(model.formula)
        for run in para.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

        _add_table(doc, model.coefficients)
        doc.add_paragraph(
            "標準化係數 β 已消去單位差異，可直接比較各變項的影響力大小；"
            "VIF 用來檢查共線性，一般以 10 為警戒線。"
        )

        doc.add_heading(f"3.{i}.3 評估指標", level=3)
        metric_table = pd.DataFrame(
            [
                {"評估方式": "樣本內", **{m: model.metrics[m] for m in METRIC_ORDER}},
                {"評估方式": "5-fold 交叉驗證", **{m: model.cv_metrics[m] for m in METRIC_ORDER}},
            ]
        )
        _add_table(doc, metric_table)

        diag_table = pd.DataFrame([{"檢定": k, "值": v} for k, v in model.diagnostics.items()])
        _add_table(doc, diag_table)
        doc.add_paragraph(
            "Durbin-Watson 接近 2 表示殘差無自我相關；Breusch-Pagan p 值大於 0.05 表示未偵測到異質變異；"
            "Jarque-Bera p 值大於 0.05 表示殘差可視為常態。"
        )

        _add_figure(doc, figures.get(f"actual_{model.target}"), f"圖 3-{i}a　{model.target} 實際值 vs 預測值")
        _add_figure(doc, figures.get(f"residual_{model.target}"), f"圖 3-{i}b　{model.target} 殘差診斷")


def _comparison_section(doc: Document, models: list[ModelResult], figures: dict[str, Path]) -> None:
    doc.add_heading("4. 模型比較", level=1)

    table = pd.DataFrame(
        [
            {
                "結果變項": m.target,
                "選入變項數": len(m.features),
                "R² (樣本內)": m.metrics["R²"],
                "Adjusted R²": m.metrics["Adjusted R²"],
                "R² (交叉驗證)": m.cv_metrics["R²"],
                "RMSE (樣本內)": m.metrics["RMSE"],
                "RMSE (交叉驗證)": m.cv_metrics["RMSE"],
            }
            for m in models
        ]
    )
    _add_table(doc, table)

    _add_figure(doc, figures.get("metrics"), "圖 4-1　六項迴歸評估指標總覽")
    _add_figure(doc, figures.get("bolasso"), "圖 4-2　Bolasso 變項選入頻率")
    _add_figure(doc, figures.get("forest"), "圖 4-3　迴歸係數森林圖")

    best = max(models, key=lambda m: m.cv_metrics["R²"] if not pd.isna(m.cv_metrics["R²"]) else -1e9)
    worst = min(models, key=lambda m: m.cv_metrics["R²"] if not pd.isna(m.cv_metrics["R²"]) else 1e9)
    if len(models) > 1:
        doc.add_paragraph(
            f"以交叉驗證 R² 來看，{best.target} 最容易被這組輸入變項解釋"
            f"（R² = {_fmt(best.cv_metrics['R²'])}），{worst.target} 最難"
            f"（R² = {_fmt(worst.cv_metrics['R²'])}）。"
        )


def _caveats_section(doc: Document, models: list[ModelResult], bolassos: dict[str, BolassoResult]) -> None:
    doc.add_heading("5. 限制與注意事項", level=1)

    items = [
        "本報告全部為相關性分析，OLS 係數描述的是變項之間的關聯強度，不能解讀為因果關係。",
        "Bolasso 的選入頻率取決於 bootstrap 抽樣，換一組亂數種子，接近門檻的變項可能進出。"
        "頻率遠高於或遠低於門檻的變項才是穩定的結論。",
        "樣本內指標必然優於交叉驗證指標。判斷模型實際預測能力請看交叉驗證那一組。",
        "Predictions.xlsx 的「誤差率 (%)」以實際值為分母，當實際值接近 0 時該欄留白 —— "
        "分母趨近於零會讓百分比失去意義。已標準化（平均 0）的資料多半屬於這種情況，"
        "此時請改看「誤差率 (佔全距 %)」。",
        "含缺失值的樣本在建模時整列刪除，各結果變項實際使用的樣本數可能不同，詳見各模型的 n。",
    ]
    fallback = [b for b in bolassos.values() if b.fallback_note]
    if fallback:
        items.append(
            "以下結果變項沒有變項達到原定選入門檻，已退讓處理，其模型結論應保守看待："
            + "、".join(b.target for b in fallback)
            + "。"
        )
    weak = [m for m in models if not pd.isna(m.cv_metrics["R²"]) and m.cv_metrics["R²"] < 0.1]
    if weak:
        items.append(
            "以下結果變項的交叉驗證 R² 低於 0.10，代表現有輸入變項幾乎無法解釋其變異："
            + "、".join(m.target for m in weak)
            + "。"
        )

    for item in items:
        doc.add_paragraph(item, style="List Bullet")
