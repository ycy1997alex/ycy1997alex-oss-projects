# -*- coding: utf-8 -*-
"""用 RegressionAnalyzer 的分析結果，產一份排版比照
「Statistical Analysis Report [Template(FILENAME)].docx」的報告。

樣板是另一套更完整的流程（GLM／GAM／Ridge 比較、1-SE rule、SEM）產出的，
RegressionAnalyzer 只做 Bolasso + OLS，所以樣板中沒有對應結果的章節（模型比較、
1-SE 選模、SEM 全章）不寫；標題塊、章節層級、表圖標題與字級色彩則照樣板。

報告的文字敘述一律由數字推導（顯著性、VIF、選入與否），換一份資料重跑，
結論會跟著變 —— 這正是「模板不動、只換資料」的用法。

用法：
    C:/Users/Alex/anaconda3/envs/stats/python.exe make_report.py Origin
    C:/Users/Alex/anaconda3/envs/stats/python.exe make_report.py Wrong
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Pt, RGBColor

HERE = Path(__file__).resolve().parent
TEMPLATE = HERE / "Statistical Analysis Report [Template(FILENAME)].docx"
DATASET = sys.argv[1] if len(sys.argv) > 1 else "Origin"
# 讀的是 RegressionAnalyzer 實際跑出來的那一次「Output [檔名] @YYYY-MM-DD HHMMSS」資料夾
OUTPUT_DIR = next((HERE / f"Output_{DATASET}").glob("Output [*"))
DATA_FILE = HERE / f"Diabetes_Data_{DATASET}.xlsx"
REPORT_FILE = HERE / f"Statistical Analysis Report [{DATA_FILE.stem}].docx"

TITLE_COLOR = RGBColor(0x1A, 0x56, 0x9E)
FILE_COLOR = RGBColor(0x1B, 0x7A, 0x55)
META_COLOR = RGBColor(0x52, 0x51, 0x4E)
BODY_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
FIG_WIDTH = Emu(5760720)


def para(doc, text, *, size=None, color=BODY_COLOR, bold=None, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    return p


def caption(doc, text):
    return para(doc, text, size=10, color=META_COLOR, center=True)


def figure(doc, relative_path, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(OUTPUT_DIR / relative_path), width=FIG_WIDTH)
    caption(doc, text)


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for cell, text in zip(t.rows[0].cells, header):
        run = cell.paragraphs[0].add_run(str(text))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    for row in rows:
        for cell, value in zip(t.add_row().cells, row):
            run = cell.paragraphs[0].add_run("" if value is None else str(value))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
    return t


def fmt(value, digits=3):
    return "" if pd.isna(value) else f"{value:.{digits}f}"


def load_results() -> dict:
    sheets = pd.read_excel(OUTPUT_DIR / "Model_Report.xlsx", sheet_name=None)
    sheets["_data"] = pd.read_excel(DATA_FILE, sheet_name="Data")
    return sheets


def build(res: dict) -> None:
    doc = Document(TEMPLATE)
    for child in list(doc.element.body):
        if not child.tag.endswith("}sectPr"):
            doc.element.body.remove(child)

    summary = res["Summary"].iloc[0]
    metrics = res["Metrics"].to_dict("records")
    coefficients = res["Coefficients"].to_dict("records")
    frequency = res["Bolasso_Frequency"].to_dict("records")
    diagnostics = res["Diagnostics"].iloc[0]
    data = res["_data"]

    # 執行時間取自輸出資料夾名的 @YYYY-MM-DD HHMMSS
    moment = datetime.strptime(OUTPUT_DIR.name.split("@")[-1].strip(), "%Y-%m-%d %H%M%S")
    when = moment.strftime("%Y年%m月%d日 - %H:%M")
    n = int(summary["樣本數 n"])
    threshold = float(summary["選入門檻"])
    # 本次實際納入的輸入變項以 Bolasso 候選清單為準（介面上可以勾掉不分析的欄位）
    inputs = [r["候選變項"] for r in frequency]
    excluded = [c for c in res["Column_Definition"]["欄位"]
                if c not in inputs and c != "target"]

    # ── 標題塊 ────────────────────────────────────────────────
    para(doc, "統計分析報告", size=24, color=TITLE_COLOR, bold=True, center=True)
    para(doc, f"檔案名稱： {DATA_FILE.stem}", size=16, color=FILE_COLOR, center=True)
    para(doc, when, size=12, color=None, center=True)
    doc.add_paragraph()
    para(doc, "分析範圍：預測模型分析（Bolasso 特徵選擇 + OLS 迴歸）", size=12, color=META_COLOR, center=True)
    para(doc, "目標變項（1）：target", size=12, color=META_COLOR, center=True)
    para(doc, f"總樣本數：{n} 筆　|　有效樣本數：{n} 筆", size=12, color=META_COLOR, center=True)

    # ── 一、預測模型分析結果 ──────────────────────────────────
    doc.add_heading("一、預測模型分析結果", level=1)
    para(
        doc,
        "本預測分析針對目標結果變項（Result Variable）建立統計建模流程，依序為："
        "(1) Bolasso 特徵選擇；(2) 以選入特徵配適普通最小平方法（OLS）迴歸；"
        "(3) 樣本內與 5-fold 交叉驗證雙軌評估；(4) 殘差診斷。"
        "所採用的估計式全屬相關性模型，係數描述的是變項之間的關聯強度，不能解讀為因果關係。",
    )

    doc.add_heading("1.1 變項基本說明", level=2)
    para(
        doc,
        f"本次分析共載入 {n} 筆原始資料，無缺失值，全數納入分析。"
        f"共納入 {len(inputs)} 個輸入變項（Input）與 1 個目標變項（Result），全部為連續型；"
        + (f"資料檔中的類別型變項 {'、'.join(excluded)} 本次未納入分析。" if excluded else "")
        + "下表為各變項的基本統計狀況。",
    )
    caption(doc, "表：連續型變項描述統計")
    table(
        doc,
        ["變項", "角色", "有效N", "缺失", "平均", "標準差", "最小", "中位數", "最大"],
        [
            [name, "Result" if name == "target" else "Input", int(data[name].count()),
             int(data[name].isna().sum()), fmt(data[name].mean()), fmt(data[name].std()),
             fmt(data[name].min()), fmt(data[name].median()), fmt(data[name].max())]
            for name in [*inputs, "target"]
        ],
    )

    doc.add_heading("1.2 分析摘要", level=2)
    para(
        doc,
        "下表列出目標變項的分析任務類型、有效樣本數、最終模型、樣本內與交叉驗證的評估分數，"
        "以及 Bolasso 選入的特徵數。完整的係數、診斷數值與選入頻率見後續小節與 Model_Report.xlsx。",
    )
    table(
        doc,
        ["目標變項", "任務類型", "有效樣本數", "最終模型", "評估指標", "樣本內分數", "交叉驗證分數",
         "RMSE（樣本內）", "RMSE（交叉驗證）", "選用特徵數"],
        [["target", "迴歸", n, "OLS", "R²", fmt(summary["R²"]), fmt(summary["R² (交叉驗證)"]),
          fmt(summary["RMSE (樣本內)"], 2), fmt(summary["RMSE (交叉驗證)"], 2),
          int(summary["選入變項數"])]],
    )

    # ── 目標變項小節 ──────────────────────────────────────────
    doc.add_heading("目標變項：target", level=2)

    doc.add_heading("特徵選擇結果（Bolasso）", level=3)
    picked = [r for r in frequency if r["是否選入"] == "是"]
    para(
        doc,
        f"Bolasso（Bootstrap LASSO）對 {len(frequency)} 個候選特徵進行 200 次 Bootstrap 重抽樣，"
        "每次以 Lasso（L1 正則化）回歸識別重要特徵，並統計各特徵的累積「Bootstrap 選擇頻率」。"
        f"懲罰強度 α = {summary['Bolasso α']:.4f}（5-fold LassoCV 於全樣本選定後固定，"
        "讓 200 次重抽樣的懲罰一致，選入頻率之間才能直接比較）。"
        f"以選擇頻率 ≥ {threshold:.0%} 為門檻，共 {len(picked)} 個特徵被選入分析"
        f"（佔 {len(picked) / len(frequency):.1%}）。選擇頻率最高的前 3 個特徵為 "
        + "、".join(f"{r['候選變項']}（{r['選入頻率']:.1%}）" for r in frequency[:3]) + "。",
    )
    caption(doc, "表：Bolasso 選擇頻率")
    table(
        doc,
        ["特徵名稱", "選擇頻率", "是否選入"],
        [[r["候選變項"], f"{r['選入頻率']:.1%}", r["是否選入"]] for r in frequency],
    )
    para(
        doc,
        "重點結論：對預測「target」而言，具穩定預測貢獻的重要變項為 "
        + "、".join(r["候選變項"] for r in picked) + "（依選擇頻率高至低排列）。",
    )
    figure(doc, "Model Analysis/01_bolasso_selection_frequency.png",
           "圖：Bolasso 特徵選擇頻率圖。橫條=Bootstrap 選中比率，虛線=門檻，深色=入選，淺色=排除。")

    doc.add_heading("迴歸模型（OLS）", level=3)
    const = next(r for r in coefficients if r["變項"] == "const")
    terms = [r for r in coefficients if r["變項"] != "const"]
    para(
        doc,
        f"target = {const['係數']:.4f} "
        + " ".join(f"{r['係數']:+.4f}·{r['變項']}" for r in terms),
    )
    para(
        doc,
        "標準化係數 β 已消去單位差異，可直接比較各變項的影響力大小；"
        "VIF 用來檢查共線性，一般以 10 為警戒線。"
        "顯著性欄以 p 值標記：*** p<0.001、** p<0.01、* p<0.05、n.s. 未達顯著。",
    )
    caption(doc, "表：OLS 迴歸係數")
    table(
        doc,
        ["變項", "係數", "標準誤", "t 值", "p 值", "CI 下界(95%)", "CI 上界(95%)",
         "標準化係數 β", "VIF", "顯著性"],
        [[r["變項"], fmt(r["係數"], 4), fmt(r["標準誤"], 4), fmt(r["t 值"]),
          f"{r['p 值']:.3e}", fmt(r["CI 下界 (95%)"], 4), fmt(r["CI 上界 (95%)"], 4),
          fmt(r["標準化係數 β"], 4), fmt(r["VIF"], 2), r["顯著性"]] for r in coefficients],
    )
    figure(doc, "Model Analysis/02_coefficient_forest.png",
           "圖：迴歸係數森林圖。點=係數估計值，橫線=95% 信賴區間；區間跨越 0 者未達統計顯著。")

    doc.add_heading("模型評估", level=3)
    in_sample = next(r for r in metrics if r["評估方式"] == "樣本內")
    cv = next(r for r in metrics if "交叉驗證" in r["評估方式"])
    para(
        doc,
        f"迴歸任務最終模型「OLS」在樣本內的 R² = {in_sample['R²']:.3f}，"
        f"即模型可解釋目標變項「target」約 {in_sample['R²']:.1%} 的變異量"
        "（R² 越接近 1 代表解釋力越強；R² < 0 表示模型預測效果劣於直接使用均值）。"
        f"5-fold 交叉驗證的 R² = {cv['R²']:.3f}，RMSE 由 {in_sample['RMSE']:.2f} 上升至 {cv['RMSE']:.2f}。"
        "樣本內指標必然偏樂觀，交叉驗證的數字才接近模型對新資料的實際表現；"
        f"兩者差距 {in_sample['R²'] - cv['R²']:.3f}，"
        + ("未見明顯過擬合。" if in_sample["R²"] - cv["R²"] < 0.05
           else "差距偏大，有過擬合疑慮。"),
    )
    caption(doc, "表：評估指標（樣本內 vs 交叉驗證）")
    table(
        doc,
        ["評估方式", "MAE", "MSE", "RMSE", "Max Error", "R²", "Adjusted R²"],
        [[r["評估方式"], fmt(r["MAE"], 2), fmt(r["MSE"], 2), fmt(r["RMSE"], 2),
          fmt(r["Max Error"], 2), fmt(r["R²"]), fmt(r["Adjusted R²"])] for r in metrics],
    )
    figure(doc, "Model Analysis/00_metrics_overview.png",
           "圖：六項迴歸評估指標總覽。每組左為樣本內、右為交叉驗證。")
    figure(doc, "Model Analysis/10_actual_vs_predicted_target.png",
           "圖：預測值 vs 實際值散佈圖。點越集中於對角線代表預測越準確。")
    para(
        doc,
        f"殘差診斷：Durbin-Watson = {diagnostics['Durbin-Watson']:.3f}（接近 2 表示殘差無自我相關）；"
        f"Breusch-Pagan p = {diagnostics['Breusch-Pagan p 值']:.4f}"
        f"（{'偵測到異質變異' if diagnostics['Breusch-Pagan p 值'] < 0.05 else '未偵測到異質變異'}）；"
        f"Jarque-Bera p = {diagnostics['Jarque-Bera p 值']:.4f}"
        f"（{'殘差可視為常態' if diagnostics['Jarque-Bera p 值'] > 0.05 else '殘差偏離常態'}）；"
        f"條件數 = {diagnostics['條件數 (Condition No.)']:.0f}；"
        f"F 統計量 = {diagnostics['F 統計量']:.2f}（p = {diagnostics['F 檢定 p 值']:.3e}），"
        "整體迴歸達統計顯著。",
    )
    figure(doc, "Model Analysis/11_residual_diagnostics_target.png",
           "圖：殘差診斷。含殘差對預測值散佈、殘差分佈與 Q-Q 圖。")

    doc.add_heading("重點關注特徵", level=3)
    freq_map = {r["候選變項"]: r["選入頻率"] for r in frequency}
    ranked = sorted(terms, key=lambda r: abs(r["標準化係數 β"]), reverse=True)
    lines = [
        f"{i}. 「{r['變項']}」：影響力 {abs(r['標準化係數 β']):.3f}，"
        f"{'正向' if r['係數'] > 0 else '負向'}（係數 {r['係數']:+.4f}），"
        f"選擇頻率 {freq_map.get(r['變項'], float('nan')):.1%}，顯著性 {r['顯著性']}"
        for i, r in enumerate(ranked, start=1)
    ]
    para(
        doc,
        "綜合前述篩選與建模結果，對「target」而言最需要優先關注的特徵如下"
        "（依標準化係數絕對值由大至小排序）：\n" + "\n".join(lines),
    )

    # ── 二、結論與建議 ────────────────────────────────────────
    doc.add_heading("二、結論與建議", level=1)
    doc.add_heading("預測模型", level=3)
    significant = [r for r in terms if r["顯著性"] != "n.s."]
    para(
        doc,
        f"• target（regression）：最終採用 OLS 模型，樣本內 R² = {in_sample['R²']:.3f}、"
        f"交叉驗證 R² = {cv['R²']:.3f}、交叉驗證 RMSE = {cv['RMSE']:.2f}。\n"
        f"• Bolasso 於 {len(frequency)} 個候選特徵中選入 {len(picked)} 個，"
        f"其中 {len(significant)} 個在 OLS 中達統計顯著（p < 0.05）。",
    )

    doc.add_heading("重點特徵", level=3)
    not_significant = [r["變項"] for r in terms if r["顯著性"] == "n.s."]
    para(
        doc,
        "依標準化係數絕對值排序，最具影響力的前三個特徵為 "
        + "、".join(f"「{r['變項']}」（β = {r['標準化係數 β']:+.3f}，{r['顯著性']}）" for r in ranked[:3])
        + "。"
        + ("其餘特徵雖被 Bolasso 選入，但在控制其他變項後未達顯著。" if not_significant
           else "進入最終模型的特徵全部達統計顯著。"),
    )

    doc.add_heading("綜合建議", level=3)
    high_vif = [r["變項"] for r in terms if r["VIF"] > 10]
    # BMI 與血糖是本資料集文獻公認的主要預測因子，兩者同時失去解釋力代表資料本身可疑
    key_features = [r for r in terms if r["變項"] in ("bmi", "glu")]
    suspicious = bool(key_features) and all(r["顯著性"] == "n.s." for r in key_features)

    bullets = []
    if high_vif:
        bullets.append(
            f"• 共線性：{'、'.join(high_vif)} 的 VIF 超過警戒線 10"
            f"（最高 {max(r['VIF'] for r in terms):.1f}），血脂類指標之間高度重疊，"
            "個別係數的正負號與大小不宜單獨解讀，建議擇一保留或改用正則化模型。"
        )
    else:
        bullets.append(
            f"• 共線性：所有進入模型的特徵 VIF 皆低於警戒線 10"
            f"（最高 {max(r['VIF'] for r in terms):.2f}），係數可各自解讀。"
        )
    if not_significant:
        bullets.append(
            f"• 選入但不顯著：{'、'.join(not_significant)} 被 Bolasso 選入卻在 OLS 中未達顯著。"
            "選入頻率高只代表該變項在重抽樣下常被留下，不等於效果顯著。"
        )
    if suspicious:
        bullets.append(
            "• 資料品質：本資料集在既有文獻中，BMI 與血糖是疾病進展最主要的預測因子；"
            "本次結果卻顯示這兩者的效果微弱且不顯著，與領域知識明顯不符。"
            f"模型整體 R² 仍有 {in_sample['R²']:.3f}，單看分數不會發現異常，"
            "建議回頭檢查資料的列對應是否正確（各欄是否來自同一筆樣本、合併時是否錯位）再行解讀。"
        )
    bullets.append("• 本報告全部為相關性分析，不能解讀為因果關係；欲確認因果需另行設計實驗或準實驗。")
    para(doc, "\n".join(bullets))

    doc.save(REPORT_FILE)
    print(f"[report] {REPORT_FILE.name}：{len(doc.paragraphs)} 段、"
          f"{len(doc.tables)} 表、{len(doc.inline_shapes)} 圖")


if __name__ == "__main__":
    build(load_results())
